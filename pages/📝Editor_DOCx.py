import streamlit as st
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def Ler_DOCx(file_path):
    DOC = Document(file_path)
    TEXTO = ''
    #TITULO = st.text_input("TÍTULO: ", DOC.paragraphs[0].text)
    Titulo = DOC.paragraphs[0].text
    n = len(DOC.paragraphs)
    for i in range(1, n):
        TEXTO+=DOC.paragraphs[i].text + '\n'
        #st.write(DOC.paragraphs[i].text)
    return Titulo, TEXTO

form = st.form('Formulário')

# Path to your .docx file
NomeArq = form.text_input("Digite nome do arquivo.DOCx e tecle ENTER: 👉", 'generactiva.docx')
check_file = os.path.isfile(NomeArq)
if check_file:
    TITULO = form.text_input("TÍTULO: ", Ler_DOCx(NomeArq)[0])
    memo = form.text_area("Conteúdo: ", Ler_DOCx(NomeArq)[1])
    if form.form_submit_button(label = '✔️ Salvar atualizações'): 
        document = Document()
        document.add_heading(TITULO, 0)            
        p = document.add_paragraph(memo)
        p.bold = True
        p.italic = True
        #p.add_run('bold').bold = True
        #p.add_run(' and some ')
        #p.add_run('italic.').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                  
        #document.add_page_break()
        document.save(NomeArq)
        #document.save('gemini03.docx')